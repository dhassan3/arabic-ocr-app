import streamlit as st
from PIL import Image
from pdf2image import convert_from_bytes
import tempfile
from pathlib import Path
import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


# -----------------------------
# Streamlit page configuration
# -----------------------------
st.set_page_config(
    page_title="Arabic Scanned PDF → Word (PaddleOCR)",
    page_icon="📄🔍",
    layout="centered"
)

st.title("🇸🇦 Scanned Arabic PDF/Image → Editable Word")
st.markdown(
    """
Upload scanned PDFs or images containing Arabic text.  
This app uses PaddleOCR to extract the text and generates an editable .docx file with RTL formatting.
"""
)


# -----------------------------
# PaddleOCR loader (cached)
# -----------------------------
@st.cache_resource(show_spinner="Loading Arabic OCR model…")
def load_paddle_ocr():
    from paddleocr import PaddleOCR
    # CPU mode for broad compatibility; change use_gpu=True if GPU available
    return PaddleOCR(use_angle_cls=True, lang="ar", use_gpu=False)


ocr = load_paddle_ocr()


# -----------------------------
# Helper functions
# -----------------------------
def pdf_to_images(file_bytes: bytes, dpi: int = 300):
    """Convert a PDF (bytes) to a list of PIL images."""
    return convert_from_bytes(file_bytes, dpi=dpi)


def run_ocr_on_image(pil_image: Image.Image) -> list[str]:
    """
    Run PaddleOCR on a PIL image and return a list of recognized text lines.
    """
    page_lines: list[str] = []

    # Use a temporary file because PaddleOCR expects a file path
    with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
        pil_image.save(tmp.name, format="JPEG")
        temp_path = tmp.name

    try:
        result = ocr.ocr(temp_path, cls=True) or []
        # result is typically a list of [ [box], [text, conf] ]
        for line in result:
            try:
                text = line[1][0]
                if isinstance(text, str) and text.strip():
                    page_lines.append(text.strip())
            except Exception:
                # robust to unexpected OCR result structures
                continue
    finally:
        # Clean up the temp image
        Path(temp_path).unlink(missing_ok=True)

    return page_lines


def add_arabic_paragraph(doc: Document, text: str):
    """
    Add a properly shaped, RTL Arabic paragraph to a Word document.
    """
    reshaped = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped)

    p = doc.add_paragraph(bidi_text)
    p.paragraph_format.right_to_left = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.paragraph_format.line_spacing = 1.15

    for run in p.runs:
        run.font.size = Pt(12)
        # Set a primary Arabic font; Word will fallback if missing
        run.font.name = "Arabic Typesetting"


# -----------------------------
# File upload UI
# -----------------------------
uploaded_file = st.file_uploader(
    "Upload scanned PDF or image",
    type=["pdf", "jpg", "jpeg", "png"]
)

if uploaded_file is None:
    st.info("Please upload a scanned PDF or an image file to get started.")
else:
    file_bytes = uploaded_file.getvalue()
    file_name = uploaded_file.name

    st.success(f"File uploaded: {file_name}")

    # Preview section
    if uploaded_file.type.startswith("image/"):
        try:
            img = Image.open(uploaded_file).convert("RGB")
            st.image(img, caption="Preview", use_column_width=True)
            images = [img]
        except Exception as e:
            st.error(f"Could not open image: {e}")
            st.stop()
    else:
        # PDF case
        with st.spinner("Converting PDF pages to images…"):
            try:
                images = pdf_to_images(file_bytes, dpi=300)
            except Exception as e:
                st.error(f"Error converting PDF to images: {e}")
                st.stop()

        st.info(f"Detected {len(images)} page(s) in the PDF.")

    # Trigger OCR + DOCX generation
    if st.button("Extract Text & Create Word File"):
        if not images:
            st.error("No pages/images were found to process.")
            st.stop()

        with st.spinner("Running Arabic OCR on all pages…"):
            doc = Document()
            doc.add_heading(f"Extracted from: {file_name}", level=1)

            all_page_texts: list[str] = []

            for page_num, page_img in enumerate(images, start=1):
                st.write(f"Processing page {page_num} of {len(images)}…")

                page_lines = run_ocr_on_image(page_img)
                page_text = "\n".join(page_lines).strip()

                if page_text:
                    all_page_texts.append(f"— Page {page_num} —\n{page_text}")

                    # For better RTL behavior, add paragraph per line
                    for line in page_lines:
                        add_arabic_paragraph(doc, line)

                    # Add a page break between pages, not after the last page
                    if page_num < len(images):
                        doc.add_page_break()
                else:
                    all_page_texts.append(f"— Page {page_num} —\n[No text detected]")

            # Save Word to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
                doc.save(tmp_docx.name)
                docx_path = tmp_docx.name

        st.success("OCR complete! Your editable Word file is ready.")

        # Download button
        base_name = Path(file_name).stem
        download_name = f"arabic_ocr_{base_name}.docx"

        with open(docx_path, "rb") as f:
            st.download_button(
                label="📥 Download Word (.docx)",
                data=f,
                file_name=download_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # Text preview
        combined_text = "\n\n".join(all_page_texts).strip()
        with st.expander("View extracted text"):
            st.text_area("Raw extracted text", value=combined_text or "No text found.", height=250)

        st.info(
            "Tip: In Word, if text direction looks off, use Right‑to‑Left paragraph direction "
            "(e.g., Ctrl + Right Shift)."
        )

st.markdown("---")
st.caption("Powered by PaddleOCR – accurate Arabic extraction ❤️")
